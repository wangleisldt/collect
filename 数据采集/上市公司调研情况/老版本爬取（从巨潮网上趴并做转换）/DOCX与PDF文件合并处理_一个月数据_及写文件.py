from 函数目录.function import file_List_Func
from 函数目录 import profile as pf
from 函数目录.function import 获取文件名前缀与后缀

from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfparser import PDFParser, PDFDocument

from 数据采集.上市公司调研情况.处理字符串获取机构名称 import 处理表格中的字符串,对PDF文件的字符串进行预处理

from 函数目录.function import checkAndCreateDir

from 函数目录.date import getYearMonthFromMonthLength

import pickle

import docx


def 处理某月所有DOCX与PDF文件(yearMonth,是否要打印清单=False):
    输出字典 = {}

    #产生需要打开的目录
    base_dir_name = "%s%s%s" % (pf.GLOBAL_PATH, pf.SEPARATOR, pf.FUNDAMENTAL_DATA)
    dirname = "%s%s%s%s%s%s%s%s" % (
    base_dir_name, pf.SEPARATOR, pf.投资者关系活动记录表, pf.SEPARATOR, pf.原始数据, pf.SEPARATOR, yearMonth[0:4], pf.SEPARATOR)

    #打开根据年月打开目录读取里面的文件
    fileList = file_List_Func(dirname)
    for filename in fileList:
        文件名前缀, 文件名后缀 = 获取文件名前缀与后缀(filename)
        # print(filename)
        #判断文件后缀与所需要处理的月份是否一直，一致则处理相应的文件
        if 文件名后缀 == 'DOCX' and 文件名前缀.split('_')[2].split('-')[1] == yearMonth[4:6]:
            print(filename)
            sum = 处理某个DOCX文件(dirname, filename,是否要打印清单)
            print(filename,sum)
            输出字典[filename] = sum
            pass

        elif 文件名后缀 == 'PDF' and 文件名前缀.split('_')[2].split('-')[1] == yearMonth[4:6]:
            print(filename)
            sum = 处理某个PDF文件(dirname, filename,是否要打印清单)
            print(filename,sum)
            输出字典[filename] = sum
            pass
        else:
            sum = 0

    #字典打印输出(输出字典)

    将字典保存成文件(输出字典,yearMonth)


def 处理某个DOCX文件(dirname,filename,是否要打印清单):
    def 打印输出(listReturn,filename,是否要打印清单):
        if 是否要打印清单:
            for element in listReturn:
                print(element, filename)

    sum = 0
    sumList = []
    开关 = False
    #doc = docx.Document(os.path.join(dirname, filename))
    doc = docx.Document(dirname+ filename)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '参与单位' in row.cells[0].text or '主办单位' in row.cells[0].text or '来访单位' in row.cells[0].text or '参加单位' in row.cells[0].text or \
                        '参加人员' in row.cells[0].text or '参与人员' in row.cells[0].text or '及人员姓名' in row.cells[0].text:
                    if table._column_count == 1:
                        开关 = True
                    else:
                        if cell.text != row.cells[0].text:
                            if len(cell.tables) > 0:
                                for t in cell.tables:
                                    for c in t._cells:
                                        #print(c.text)
                                        cell.text = c.text
                                        listReturn = 处理表格中的字符串(cell.text)
                                        sumList = sumList + listReturn
                                        打印输出(listReturn, filename, 是否要打印清单)
                            else:
                                listReturn = 处理表格中的字符串(cell.text)
                                sumList = sumList + listReturn
                                打印输出(listReturn, filename,是否要打印清单)
                elif 开关 == True:
                    listReturn = 处理表格中的字符串(cell.text)
                    sumList = sumList + listReturn
                    打印输出(listReturn, filename,是否要打印清单)
                    开关 = False

    sum = len(list(set(sumList)))#对结果进行去重
    if sum is None:
        sum = 0
    return sum



def 处理某个PDF文件(dirname,filename,是否要打印清单):
    def 匹配是否有结束字符串(year,results):
        需要匹配的字符串 = [year, str(int(year) - 1), '公司会议室', '接待人员']
        for element in 需要匹配的字符串:
            if element in results:
                return True
        return False

    try:
        有起始的字符串 = False
        有结束的字符串 = False
        year = filename.split('_')[2].split('-')[0]
        sum = 0
        sumList = []




        fp = open(dirname+filename, 'rb')  # rb以二进制读模式打开本地pdf文件
        # request = Request(url=_path, headers={'User-Agent': random.choice(user_agent)})  # 随机从user_agent列表中抽取一个元素
        # fp = urlopen(request) #打开在线PDF文档

        # 用文件对象来创建一个pdf文档分析器
        praser_pdf = PDFParser(fp)

        # 创建一个PDF文档
        doc = PDFDocument()

        # 连接分析器 与文档对象
        praser_pdf.set_document(doc)
        doc.set_parser(praser_pdf)

        # 提供初始化密码doc.initialize("123456")
        # 如果没有密码 就创建一个空的字符串
        doc.initialize()

        # 检测文档是否提供txt转换，不提供就忽略
        if not doc.is_extractable:
            pass
            # raise PDFTextExtractionNotAllowed
        else:
            # 创建PDf资源管理器 来管理共享资源
            rsrcmgr = PDFResourceManager()

            # 创建一个PDF参数分析器
            laparams = LAParams()

            # 创建聚合器
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)

            # 创建一个PDF页面解释器对象
            interpreter = PDFPageInterpreter(rsrcmgr, device)

            # 循环遍历列表，每次处理一页的内容
            # doc.get_pages() 获取page列表
            for page in doc.get_pages():
                # 使用页面解释器来读取
                interpreter.process_page(page)

                # 使用聚合器获取内容
                layout = device.get_result()

                # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等 想要获取文本就获得对象的text属性，
                for out in layout:
                    # 判断是否含有get_text()方法，图片之类的就没有
                    # if hasattr(out,"get_text"):
                    if isinstance(out, LTTextBoxHorizontal):
                        results = out.get_text().strip()
                        #print(results)

                        if '现场参观' in results and not 有结束的字符串:
                            # print("results: " + results)
                            有起始的字符串 = True
                        elif 有起始的字符串 and not 有结束的字符串:

                            if 匹配是否有结束字符串(year,results):
                            #if year in results or str(int(year) - 1) in results or '公司会议室' in results or '接待人员' in results:
                                有结束的字符串 = True
                            else:
                                #print(results)
                                results = 对PDF文件的字符串进行预处理(results)
                                #print(results)
                                returnList = 处理表格中的字符串(results)
                                if 是否要打印清单:
                                    for element in returnList:
                                        print(element,filename)
                                sumList = sumList+returnList

        sum = len(list(set(sumList)))  # 对结果进行去重
        if sum is None:
            sum = 0
        return sum
    except:
        print("文件处理出错！")
        return 0

def 字典打印输出(dict):
    for key in dict:
        print(key,dict[key])

def 将字典保存成文件(dict,yearMonth):
    # 产生需要打开的目录
    base_dir_name = "%s%s%s" % (pf.GLOBAL_PATH, pf.SEPARATOR, pf.FUNDAMENTAL_DATA)
    dirname = "%s%s%s%s%s%s" % (
        base_dir_name, pf.SEPARATOR, pf.投资者关系活动记录表, pf.SEPARATOR, pf.汇总数据, pf.SEPARATOR)
    checkAndCreateDir(dirname)

    filename = '%s%s%s' % (dirname, yearMonth, pf.PklFile)

    output = open(filename, 'wb')
    pickle.dump(dict, output)
    output.close()

def 处理某月所有DOCX与PDF文件接口(startDate,monthLength):
    for i in range(0, int(monthLength)):
        yearMonth = getYearMonthFromMonthLength(startDate, i)
        处理某月所有DOCX与PDF文件(yearMonth, 是否要打印清单=False)

if __name__ == '__main__':
    处理某月所有DOCX与PDF文件接口('202001', 6)

